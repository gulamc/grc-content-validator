"""DOCX document parser for extracting content from Word documents."""

from docx import Document
from typing import Dict, List, Any


def parse_docx(file_path: str) -> Dict[str, Any]:
    """
    Parse a DOCX file and extract structured content.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Dictionary containing:
        - full_text: Complete text content
        - paragraphs: List of paragraph texts
        - structure: Document structure information
    """
    try:
        doc = Document(file_path)

        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # Only include non-empty paragraphs
                paragraphs.append(text)

        # Build full text
        full_text = '\n'.join(paragraphs)

        # Extract structure information
        structure = {
            'total_paragraphs': len(paragraphs),
            'total_characters': len(full_text),
            'total_words': len(full_text.split()),
            'has_tables': len(doc.tables) > 0,
            'table_count': len(doc.tables),
        }

        return {
            'full_text': full_text,
            'paragraphs': paragraphs,
            'structure': structure
        }

    except Exception as e:
        raise Exception(f"Error parsing DOCX file: {str(e)}")


def get_text_from_docx(file_path: str) -> str:
    """
    Extract just the text content from a DOCX file.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Full text content as string
    """
    result = parse_docx(file_path)
    return result['full_text']
